import io
import uuid
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required, user_passes_test
from django.contrib import messages
from django.contrib.auth import get_user_model
from django.http import HttpResponse
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
import openpyxl

from .models import StaffProfile, InventoryItem, StaffReport, Sale, AuditLog, CartItem, Order, OrderItem
from .forms import StaffProfileForm, InventoryItemForm, SaleForm

User = get_user_model()

def is_staff_member(user):
    return user.is_active and (user.is_staff or user.is_superuser)

def log_action(user, action, details=""):
    AuditLog.objects.create(user=user, action=action, details=details)

@login_required
def staff_list(request):
    users = User.objects.all().order_by("-date_joined")
    return render(request, "staff/staff_list.html", {"staff_members": users})

@login_required
@user_passes_test(is_staff_member)
def staff_create(request):
    if request.method == "POST":
        form = StaffProfileForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, "Staff profile created successfully.")
            return redirect("staff_list")
    else:
        form = StaffProfileForm()
    return render(request, "staff/staff_form.html", {"form": form})

@login_required
def inventory_list(request):
    try:
        user_dept = getattr(request.user, "staff_profile", None)
        if user_dept and user_dept.department and not request.user.is_superuser:
            items = InventoryItem.objects.filter(department=user_dept.department)
        else:
            items = InventoryItem.objects.all()
    except Exception:
        items = InventoryItem.objects.all()
        
    return render(request, "staff/inventory_list.html", {"items": items})

@login_required
@user_passes_test(is_staff_member)
def inventory_create(request):
    if request.method == "POST":
        form = InventoryItemForm(request.POST)
        if form.is_valid():
            item = form.save()
            log_action(request.user, "Created Inventory Item", f"Item: {item.item_name} (SKU: {item.sku})")
            messages.success(request, "Inventory item added successfully.")
            return redirect("inventory_list")
    else:
        form = InventoryItemForm()
    return render(request, "staff/inventory_form.html", {"form": form, "title": "Add New Inventory Item"})

@login_required
@user_passes_test(is_staff_member)
def record_sale(request):
    if request.method == "POST":
        form = SaleForm(request.POST)
        if form.is_valid():
            sale = form.save(commit=False)
            item = sale.item
            
            if item.quantity < sale.quantity_sold:
                messages.error(request, f"Insufficient stock! Only {item.quantity} units available.")
            else:
                item.quantity -= sale.quantity_sold
                item.save()
                
                sale.performed_by = request.user
                sale.total_price = sale.quantity_sold * item.unit_price
                sale.save()

                log_action(request.user, "Recorded Sale", f"Sold {sale.quantity_sold}x {item.item_name} for ${sale.total_price}")
                messages.success(request, "Sale recorded and inventory stock updated.")
                return redirect("sales_list")
    else:
        form = SaleForm()
    return render(request, "staff/sale_form.html", {"form": form})

@login_required
def sales_list(request):
    sales = Sale.objects.select_related("item", "performed_by").all().order_by("-sold_at")
    return render(request, "staff/sales_list.html", {"sales": sales})

@login_required
@user_passes_test(lambda u: u.is_superuser)
def audit_log_list(request):
    logs = AuditLog.objects.select_related("user").all().order_by("-timestamp")[:100]
    return render(request, "staff/audit_logs.html", {"logs": logs})

@login_required
def create_report(request):
    if request.method == "POST":
        title = request.POST.get("title")
        content = request.POST.get("content")
        equation_text = request.POST.get("equation", "").strip()
        uploaded_image = request.FILES.get("report_image")
        
        report = StaffReport.objects.create(title=title, author=request.user, content=content)
        log_action(request.user, "Generated Report", f"Title: {title}")
        
        doc = Document()
        doc.add_heading(title, 0)
        
        meta_table = doc.add_table(rows=2, cols=2)
        meta_table.style = "Table Grid"
        meta_table.cell(0, 0).text = "Author:"
        meta_table.cell(0, 1).text = request.user.username
        meta_table.cell(1, 0).text = "Generated Date:"
        meta_table.cell(1, 1).text = report.created_at.strftime("%Y-%m-%d %H:%M")

        doc.add_paragraph()
        doc.add_heading("Report Content", level=1)
        doc.add_paragraph(content)

        if equation_text:
            doc.add_heading("Mathematical Formulas", level=1)
            eq_p = doc.add_paragraph()
            eq_p.add_run("Equation: ")
            m_math = OxmlElement("m:oMathPara")
            m_math_inner = OxmlElement("m:oMath")
            m_r = OxmlElement("m:r")
            m_t = OxmlElement("m:t")
            m_t.text = equation_text
            m_r.append(m_t)
            m_math_inner.append(m_r)
            m_math.append(m_math_inner)
            eq_p._p.append(m_math)

        if uploaded_image:
            doc.add_heading("Attached Figures", level=1)
            image_stream = io.BytesIO(uploaded_image.read())
            doc.add_picture(image_stream, width=Inches(5.0))

        response = HttpResponse(content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        response["Content-Disposition"] = f'attachment; filename="{title}.docx"'
        doc.save(response)
        return response

    return render(request, "staff/create_report.html")

@login_required
def export_inventory_excel(request):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Inventory Summary"
    ws.append(["ID", "Item Name", "SKU", "Quantity", "Unit Price ($)"])

    for item in InventoryItem.objects.all():
        ws.append([item.id, item.item_name, item.sku, item.quantity, float(item.unit_price)])

    response = HttpResponse(content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    response["Content-Disposition"] = 'attachment; filename="Inventory_Report.xlsx"'
    wb.save(response)
    return response

# --- E-Commerce & Storefront Views ---

def store_home(request):
    try:
        items = InventoryItem.objects.all()
    except Exception:
        items = []
    return render(request, "store/store_home.html", {"items": items})

@login_required
def add_to_cart(request, item_id):
    item = get_object_or_404(InventoryItem, id=item_id)
    if item.quantity < 1:
        messages.error(request, f"{item.item_name} is out of stock.")
        return redirect("store_home")

    cart_item, created = CartItem.objects.get_or_create(user=request.user, item=item)
    if not created:
        if cart_item.quantity + 1 > item.quantity:
            messages.warning(request, f"Cannot add more. Only {item.quantity} available in stock.")
            return redirect("cart_view")
        cart_item.quantity += 1
        cart_item.save()
    else:
        cart_item.quantity = 1
        cart_item.save()

    messages.success(request, f"Added {item.item_name} to your cart.")
    return redirect("cart_view")

@login_required
def cart_view(request):
    try:
        items = CartItem.objects.filter(user=request.user).select_related("item")
        total = sum(item.subtotal() for item in items)
    except Exception:
        items = []
        total = 0
    return render(request, "store/cart.html", {"cart_items": items, "total": total})

@login_required
def remove_from_cart(request, item_id):
    CartItem.objects.filter(user=request.user, item_id=item_id).delete()
    messages.info(request, "Item removed from shopping cart.")
    return redirect("cart_view")

@login_required
def checkout(request):
    try:
        cart_items = CartItem.objects.filter(user=request.user).select_related("item")
    except Exception:
        cart_items = CartItem.objects.none()

    if not cart_items.exists():
        messages.error(request, "Your shopping cart is empty.")
        return redirect("store_home")

    total = sum(item.subtotal() for item in cart_items)

    if request.method == "POST":
        address = request.POST.get("shipping_address")
        phone = request.POST.get("contact_phone")
        order_num = f"ORD-{uuid.uuid4().hex[:8].upper()}"

        order = Order.objects.create(
            customer=request.user,
            order_number=order_num,
            shipping_address=address,
            contact_phone=phone,
            total_amount=total,
            status="Pending"
        )

        for ci in cart_items:
            OrderItem.objects.create(
                order=order,
                item=ci.item,
                quantity=ci.quantity,
                unit_price=ci.item.unit_price
            )
            ci.item.quantity -= ci.quantity
            ci.item.save()

            Sale.objects.create(
                item=ci.item,
                performed_by=request.user,
                quantity_sold=ci.quantity,
                total_price=ci.subtotal()
            )

        cart_items.delete()
        log_action(request.user, "Placed Online Order", f"Order #{order.order_number} Total: ${total}")
        messages.success(request, f"Order #{order.order_number} placed successfully!")
        return redirect("customer_dashboard")

    return render(request, "store/checkout.html", {"cart_items": cart_items, "total": total})

@login_required
def customer_dashboard(request):
    try:
        orders = Order.objects.filter(customer=request.user).prefetch_related("items__item").order_by("-created_at")
    except Exception:
        orders = []
    return render(request, "store/customer_dashboard.html", {"orders": orders})